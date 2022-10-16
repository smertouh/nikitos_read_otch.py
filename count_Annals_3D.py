import docx #pip install python-docx
import glob
import os
from pathlib import Path
import fnmatch
import subprocess

import aspose.words as aw #pip install aspose-words
#p=Path()
#docxlist0=glob.glob("*.docx")
#print(list(p.glob("**/*")))
#((os.walk('').filenames))
flag=0
flag2=0
Dict_arend={}

fpath0="C://Users//Home//PycharmProjects//nikitos_read_otch"#'C://Users//Никитос//PycharmProjects//xlsFunc1/'

def read_дата_поступления(par, vec=[]):
    s0=par.text.lower()
    if not "вх. №" in par.text and not "вх.№" in s0:
        if ("вх." in par.text) and "от " in s0:
            try:
                s1= s0.rsplit("от ",1)[0].rsplit("вх.")[1]
                if "г." in s0:
                    s2 = s0.rsplit("от ",1)[-1].rsplit(" г.")[0]
                else:
                    s2 = s0.rsplit("от ",1)[-1]

                print(1)
                return [s1, s2]
            except:
                return vec

    if ("вх. №" in par.text or "вх.№" in s0) and "от " in s0:
        try:
            s1= s0.rsplit("от ",1)[0].rsplit("вх.")[1]
            if "г." in s0:
                s2 = s0.rsplit("от ",1)[-1].rsplit(" г.")[0]
            else:
                s2 = s0.rsplit("от ",1)[-1]

            print(1)
            return [s1, s2]
        except:
            return vec

    return vec
def read_дата_составления_1(par):
    try:
        s0 = par.text.lower()
    except:
        s0 = par
    global flag
    if flag==1:
        #print(s0)

        if "от " in s0:
            if "№" in  s0:
                try:
                    if len(s0.rsplit("от ")[0].rsplit("№"))==2:
                        s1=s0.rsplit("от ",1)[0].rsplit("№")[1]
                    elif len(s0.rsplit("от ")[0].rsplit("№№")) == 2:
                        s1 = s0.rsplit("от ", 1)[0].rsplit("№№")[1]
                    if "г." in s0:
                        s2=s0.rsplit("г.", 1)[0].rsplit("от ")[1]
                    else:
                        s2=s0.rsplit("от ")[1]
                    flag = 0
                    return [s1,s2]
                except:
                    flag = 0
                    return []
            else:

                s2=s0.rsplit("от ", 1)[0]
                if "отчет" in s2:
                    s1 = s0.rsplit("от ", 1)[0].rsplit("отчет")[1]
                else:
                    s1=""

                if "г." in s0:
                    s2 = s0.rsplit("г.", 1)[0].rsplit("от ")[1]
                else:
                    s2 = s0.rsplit("от ")[1]
                flag = 0
                if s1!="":
                    return [s1, s2]
                else:
                    return []
        flag = 0
    if "мотивированное мнение" in s0:
        flag=1

    return []
def read_дата_составления_2(par):
    try:
        s0 = par.text.lower()
    except:
        s0 = par
    global flag2
    if "22052" in s0:
        print(121412)
    if flag2==1:
        #print(s0)
        flag=0
        if "от " in s0:# and "г." in s0:
            s1=s0.rsplit("от ",1)[0].rsplit("№")[1]
            if "г." in s0:
                s2 = s0.rsplit("г.", 1)[0].rsplit("от ")[1]
            else:
                s2 = s0.rsplit("от ")[1]
            return [s1,s2]
    if "мотивированное мнение" in s0:
        flag=1
    return []
def read_дата_составления_3(cell,svec):
    s0 = cell.lower()
    if "\n" in s0:
        snew=s0.split("\n")
    else:
        snew=[s0]
    for s0 in snew:
        s1 = svec[0]
        s2 = svec[1]
        s3 = svec[2]
        s4 = svec[3]
        s5 = svec[4]
        if s0 == "Отчет 0522АН01 от 05.05.2022г.Отчет 0522АН01 от 05.05.2022г.".lower():
            print(1)
        if "кв.м" in s0:
            s3 = read_площадь(s0)
        if "кв. м" in parag:  # 2022-09-04
            aread = read_площадь1(parag)
        if "руб" in s0:
            s4 = read_стоимость(s0)
        if "объект оценки" in s0 or "объект аренды" in s0:
            s5 = read_объект_оценки(s0)
        if "вариант наиболее эффективного использования" in s0:
            global ef_use
            ef_use = (s0.split('– ', 1)[-1].split('- ')[-1].split('— ')[-1])
        if "по состоянию на" in s0:
            global по_состоянию_на
            по_состоянию_на = read_по_состоянию_на(s0)

        if "отч" in s0 and "от " in s0:
            if "№" in s0:
                s1 = s0.rsplit("от ", 1)[0].rsplit("№")[1]
                s2 = s0.rsplit("г", 1)[0].rsplit("от ")[1]
            else:
                if "отчет " in s0:
                    s1 = s0.rsplit("от ", 1)[0].rsplit("отчет ")[1]
                    s2 = s0.rsplit("г", 1)[0].rsplit("от ")[1]
            #return [s1, s2, s3, s4, s5]
    return [s1, s2, s3, s4, s5]

def read_по_состоянию_на(par):
    s=par.split("по состоянию на")
    s2=s[1]
    s2.strip()
    i=0
    si=s2[i]
    while si in ",0123456789-. :"or si==' ':
        i+=1
        if i<len(s2):
            si = s2[i]
        else:
            break
    imax=i
    snew1=""
    for i in range(1,imax):

        si=s2[i]
        if si==",":
            si="."
        if si==' ' or si==" ":
            si=''
        snew1 += si
    #print(snew1)
    if len(snew1)!=0:
       return (snew1)+"#R"
    return "НЕТ ДАТЫ ОЦЕНКИ"
def read_дата_составления_4(par):
    for runcircle in par.runs:
        if runcircle.bold == 1:
            parag = par.text.lower()
            parag1 = ""
            for s in parag:
                if s != ' ':
                    parag1 += s
                else:
                    parag1 += " "
            par = parag1
            try:
                s0 = par.text.lower()
            except:
                s0 = par
            if "от " in s0 and ("отчет" in s0 or "отчёт" in s0) :
                if "№" in s0:
                    s1=s0.rsplit("№",1)[1].rsplit("от ",1)[0]
                    if "г." in s0:
                        print(s0)
                        s2=s0.rsplit("г.", 1)[0].rsplit("от ")[1]
                    else:
                        s2=s0.rsplit("от ")[1]
                    return [s1,s2]
                else:
                    s1 = s0.rsplit("отчет", 1)[1].rsplit("от ", 1)[0]
                    if "г." in s0:
                        print(s0)
                        s2 = s0.rsplit("г.", 1)[0].rsplit("от ")[1]
                    else:
                        s2 = s0.rsplit("от ")[1]
                    return [s1, s2]
            return []
    return []


def read_объект_оценки(par):
    return (par)

def read_дата_оценки(s0):
    pass
def read_площадь(s0):
    print(1)
    stot=0
    s1 = s0.rsplit("кв.м")
    for s2 in s1:
        if s2==s1[-1]:
            break
        s2.strip()
        i=-1
        si=s2[i]
        while si in ",0123456789 "or si==' ':
            i-=1
            si = s2[i]
        i+=1
        snew1=""
        while i<0:

            si=s2[i]
            if si==",":
                si="."
            if si==' ' or si==" ":
                si=''
            snew1 += si
            i+=1
        #print(snew1)
        if len(snew1)!=0:
            try:
                stot=max(float(snew1),stot)
            except:
                stot+=0
            print(stot)
    return (stot)

def read_площадь1(s0):
    print(1)
    stot=0
    s1 = s0.rsplit("кв. м")
    for s2 in s1:
        if s2==s1[-1]:
            break
        s2.strip()
        i=-1
        si=s2[i]
        while si in ",0123456789 "or si==' ':
            i-=1
            si = s2[i]
        i+=1
        snew1=""
        while i<0:

            si=s2[i]
            if si==",":
                si="."
            if si==' ' or si==" ":
                si=''
            snew1 += si
            i+=1
        #print(snew1)
        if len(snew1)!=0:
            try:
                stot=max(float(snew1),stot)
            except:
                stot+=0
            print(stot)
    return (stot)

def read_стоимость(s0):
    print(1)
    stot = 0
    s1 = s0.rsplit("руб")
    for s2 in s1:
        s2.strip()
        i = -1
        si = s2[i]
        while si in ",0123456789  "or si==' ':
            i -= 1
            if -len(s2)<i:
                si = s2[i]
            else:
                si="loh"
                print(1)
        i += 1
        snew1 = ""
        while i < 0:

            si = s2[i]
            if si == ",":
                si = "."
            if si==' ' or si==" ":
                si=''
            snew1 += si
            i += 1
        # print(snew1)
        if len(snew1) != 0:
            try:
                stot += float(snew1)
            except:
                stot += 0
            print(stot)
    return (stot)
def read_номер_отчёта(par):
    pass




pathes=glob.glob("**/**/*.docx")
"""
for s in os.walk('C://Users//Никитос//PycharmProjects//xlsFunc1/*.docx'):
    if 'docx' in s[2]:
        pass
    for sss in s[2]:
        if '.docx' in sss:
            for ss in s[1]:
                pathes.append(s[0] + ss)
    #pathes.append(s)
"""
matches = []




#2022-09-02 все док в докх
for root, dirnames, filenames in os.walk(fpath0):
    for filename in fnmatch.filter(filenames, '*.doc'):

        if not "~" in filename:
            print(root,"\t",filename)
            doc = aw.Document(root+"\\"+filename)
            try:
                doc.save(root+"\\"+filename + "x")
            except:
                print("тьфу тьфу")
            if not root in matches:
                matches.append(root)




for root, dirnames, filenames in os.walk(fpath0):
    for filename in fnmatch.filter(filenames, '*.docx'):
        if 'Письмо' in filename:
            pass
        else:
            if not root in matches:
                matches.append(root)
"""
#2022-08-31
for root, dirnames, filenames in os.walk(fpath0):
    for filename in fnmatch.filter(filenames, '*.doc'):
        if 'Письмо' in filename:
            pass
        else:
            if not root in matches:
                doc = aw.Document(root+"\\"+filename)
                doc.save(root+"\\"+filename + "x")
                matches.append(root)
"""
"""
for filename in fnmatch.filter(filenames, '*.doc'):
    matches.append(root)
"""
pathes = matches#['C://Users//Никитос//PycharmProjects//xlsFunc1//sdf//2019//ГПНТБ СО РАН//2019.03.27//№ 3076']#matches
#"C://Users//Никитос//PycharmProjects//xlsFunc1/"
with open("name_of_all_files.txt", "w") as file2:
    number_files = 0
    number_zagolovok=1
    zagolovok = ""
    for onepath in pathes:
        print(os.path.isdir(onepath))
        #doclist = glob.glob(onepath+"//*.doc")
        """
        for doc in doclist:
            print("открываю файл" + doc)
            subprocess.call("C://Program Files//LibreOffice//program//soffice.exe --headless --convert-to docx " + doc)
        """
        """
        for filename in os.listdir(onepath):
            if filename.endswith('.doc'):
                print (filename)
                #os.join()
                subprocess.call(['C://Program Files//LibreOffice//program//soffice.exe ', '--headless', '--convert-to', 'docx', os.path.join(onepath,'*.doc')])#,'--outdir','c:\\tmp','с:\\tmp\\1.docx'])
        """
        #file2.writelines("шапка \n")

        for f in os.scandir(onepath):
            #if  or f.is_file() and f.path.split('.')[-1].lower() == 'doc' :


            if f.is_file() and f.path.split('.')[-1].lower() == 'docx':
                number_files+=1
                try:
                    if zagolovok != "":
                        number_zagolovok += 1
                        zagolovok = ""
                except:
                    pass

                filepath=f.path
                # with open(f.path, 'r') as docxfile:
                if 'Письмо' and "Документ" in f.name:
                    break
                if not 'отч' in f.name:
                    break
                # печатаем название файла
                file2.writelines(f.name+"\t")
                date_ot = []
                svec=["","","","","",""]

                if not '~$' in f.name:
                    aread="-"
                    price="-"
                    obj_arr="-"
                    doc = docx.Document(f.path)
                    #печатаем путь
                    file2.writelines("\\"+f.path.split('\\',1)[-1]+"\t")
                    # ssPar=docxfile#.text #'.paragraph.Paragraph()
                    # print(ssPar)
                    #file2.writelines('\n')





                    vivod="НЕТ ВЫВОДА"
                    ef_use="НЕТ ЭФ. ИСП."
                    по_состоянию_на="НЕТ ПО СОСТ НА"
                    date_ekspert="НЕТ ДАТЫ ЭКСПЕРТИЗЫ"
                    for par in doc.paragraphs:
                        if par.alignment == 1:
                            for runcircle in par.runs:
                                if runcircle.bold == 1:
                                    zagolovok += par.text
                                    break
                        parag=par.text.lower()
                        parag1=""
                        for s in parag:
                            if s!=' ':
                                parag1+=s
                            else:
                                parag1+=" "
                        parag=parag1


                        if 'вариант наиболее эффективного использования' in parag:
                            #file2.writelines('\t')
                            ef_use=(parag.split('– ',1)[-1].split('- ')[-1].split('— ')[-1])
                        if 'вывод, что' in parag:
                            s=parag
                            vivod=(s.split('вывод, что ', 1)[-1].rsplit('законодательству', 1)[0]) #s - parag
                        if "руб" in parag:
                            price=read_стоимость(parag)
                        if "кв.м" in parag:
                            aread=read_площадь(parag)
                        if "кв. м" in parag:#2022-09-04
                            aread=read_площадь1(parag)
                        if "объект оценки" in parag or "объект аренды" in parag:
                            obj_arr = read_объект_оценки(parag)
                        if "дата проведения экспертизы" in parag:
                            if ":" in parag.rsplit("дата проведения экспертизы")[1]:
                                date_ekspert=parag.rsplit("дата проведения экспертизы")[1].split(":",1)[1]+"#U"
                            else:
                                print("Вы меня НЕ обманете!")
                                date_ekspert = parag.rsplit("дата проведения экспертизы")[1]+"#U"

                        elif "дата рассмотрения отчет" in parag:
                            if ":" in parag.rsplit("дата рассмотрения отчет")[1]:
                                date_ekspert = parag.rsplit("дата рассмотрения отчет")[1].split(":",1)[1]+"#U"
                            else:
                                print("Вы меня НЕ обманете! (не в этот раз!)")
                                date_ekspert = parag.rsplit("дата рассмотрения отчет")[1]+"#U"
                        if "по состоянию на" in parag:
                            по_состоянию_на=read_по_состоянию_на(parag)
                            print(по_состоянию_на)
                        if date_ekspert=="":
                            print(1)
                        date_ot=read_дата_поступления(par,date_ot)
                        date_ot1=read_дата_составления_1(parag)
                        read_дата_составления_2(parag)
                        if date_ot1!=[]:
                            try:
                                svec=[date_ot1[0],date_ot1[1],aread,price,obj_arr]
                            except:
                                pass
                        else:
                            date_ot1=[svec[0],svec[1]]
                            try:
                                svec=[date_ot1[0],date_ot1[1],aread,price,obj_arr]
                            except:
                                pass
                    if svec[3]==svec[4]==svec[2]:
                        pass
                #для списка

                #vivod = "НЕТ ВЫВОДА"
                for par in doc.paragraphs:
                    parag = par.text.lower()
                    parag1 = ""
                    for s in parag:
                        if s != ' ':
                            parag1 += s
                        else:
                            parag1 += " "
                    parag = parag1
                    if 'вариант наиболее эффективного использования' in parag:
                        # file2.writelines('\t')
                        ef_use = (parag.split('– ', 1)[-1].split('- ')[-1].split('— ')[-1])
                    if 'вывод, что' in parag:
                        s = parag
                        vivod = (s.split('вывод, что ', 1)[-1].rsplit('законодательству', 1)[0]) #s - parag
                    if "руб" in parag:
                        price = read_стоимость(parag)
                    if "кв.м" in parag:
                        aread = read_площадь(parag)
                    if "кв. м" in parag:  # 2022-09-04
                        aread = read_площадь1(parag)
                    if "объект оценки" in parag or "объект аренды" in parag:
                        obj_arr = read_объект_оценки(parag)
                    date_ot1 = read_дата_составления_4(par)
                    if date_ot1 != []:
                        try:
                            svec = [date_ot1[0], date_ot1[1], aread, price, obj_arr]
                        except:
                            pass
                    else:
                        date_ot1 = [svec[0], svec[1]]
                        try:
                            svec = [date_ot1[0], date_ot1[1], aread, price, obj_arr]
                        except:
                            pass
                    if svec[0]!="" and svec[1]!="":
                        svec2 = svec
                        if not vivod in svec2:
                            svec2.append(vivod)
                        if not filepath in svec2:
                            svec2.append(filepath)
                        if svec2 != [] and date_ot != []:
                            if not date_ot[0] in svec2:
                                svec2.append(date_ot[0])
                        if not ef_use in svec2:
                            svec2.append(ef_use)
                        if not date_ekspert in svec2:
                            svec2.append(date_ekspert)
                        if date_ot!=[]:
                            if not date_ot[1] + "#V" in svec2:
                                svec2.append(date_ot[1] + "#V")
                        if not по_состоянию_на in svec2:
                            svec2.append(по_состоянию_на)
                        if svec!=[] and date_ot!=[]:
                            Dict_arend.update({(svec[0] + "_" + date_ot[0]): svec2})



                data_tables = {i: None for i in range(len(doc.tables))}
                #svec = ["-", "-", "-", "-", "-"]
                for i, table in enumerate(doc.tables):
                    #print('\nДанные таблицы №', i)
                    # создаем список строк для таблицы `i` (пока пустые)
                    svec=["-","-","-","-","-"]
                    data_tables[i] = [[] for _ in range(len(table.rows))]
                    # проходимся по строкам таблицы `i`
                    for j, row in enumerate(table.rows):
                        # проходимся по ячейкам таблицы `i` и строки `j`
                        for cell in row.cells:
                            # добавляем значение ячейки в соответствующий
                            # список, созданного словаря под данные таблиц
                            data_tables[i][j].append(cell.text)
                            svec=read_дата_составления_3(cell.text,svec)
                            svec2=svec
                            if not vivod in svec2:
                                svec2.append(vivod)
                            if not filepath in svec2:
                                svec2.append(filepath)
                            if svec2 != [] and date_ot != []:
                                if not date_ot[0] in svec2:
                                    svec2.append(date_ot[0])
                            if not ef_use in svec2:
                                svec2.append(ef_use)
                            if not date_ekspert in svec2:
                                svec2.append(date_ekspert)
                            if date_ot != []:
                                if not date_ot[1]+"#V" in svec2:
                                    svec2.append(date_ot[1]+"#V")
                            if not по_состоянию_на in svec2:
                                svec2.append(по_состоянию_на)
                            if svec!=[] and date_ot!=[]:
                                Dict_arend.update({(svec[0] + "_" + date_ot[0]): svec2})
                    # смотрим извлеченные данные
                    # (по строкам) для таблицы `i`
                    #print(data_tables[i])
                    #print('\n')

                svec2 = svec
                if not vivod in svec2:
                    svec2.append(vivod)
                if not filepath in svec2:
                    svec2.append(filepath)
                if svec2 != [] and date_ot != []:
                    if not date_ot[0] in svec2:
                        svec2.append(date_ot[0])
                if not ef_use in svec2:
                    svec2.append(ef_use)
                if not date_ekspert in svec2:
                    svec2.append(date_ekspert)
                if date_ot != []:
                    if not date_ot[1] + "#V" in svec2:
                        svec2.append(date_ot[1] + "#V")
                if not по_состоянию_на in svec2:
                    svec2.append(по_состоянию_на)
                if svec!=[] and date_ot!=[]:
                    Dict_arend.update({(svec[0] + "_" + date_ot[0]): svec2})

                s=''
                doc=''
                par=''
                runcircle=''
                print(1)



print(Dict_arend)
print("количество отчётов \t",len(Dict_arend),"\t количество файлов \t",number_files,"\t количество заголовков \t",number_zagolovok)

import csv
with open("result.csv", 'w') as csvfile:
    spamwriter = csv.writer(csvfile, delimiter='\t')
    for key in Dict_arend:
        stringg=[key]
        for s in Dict_arend[key]:
            try:
                s=(float(s))
                s = str(s)
                s=s.replace(".", ",")
            except:
                pass
            if "#U" in s:
                s=s.split("#U")[0]
            if "#R" in s:
                s=s.split("#R")[0]
            if "#V" in s:
                s = s.split("#V")[0]
            elif "\n" in s:
                ss=s.split("\n")
                snew=""
                for sss in ss:
                    snew+=sss+" "
                s=snew
            stringg.append(s)
        try:
            spamwriter.writerow(stringg)
        except:
            spamwriter.writerow(key)

""""
print(f.path)
word = win32com.client.Dispatch("Word.Application")
word.visible = False
wb = word.Documents.Open("C:\\Users\\Никитос\\PycharmProjects\\xlsFunc1\\sdf\\2021\\банкроты\\отрицательные\\ФГУП Опытный завод\\54-724\\1 отч., Закл. ФГУП Опытный завод, БС.doc")
doc = word.ActiveDocument
print(doc.Range().Text)
"""